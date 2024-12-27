import os
import tempfile
import json
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
from langchain_community.document_loaders import PyPDFLoader
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from llama_cpp import Llama

app = FastAPI()

# Define the Llama model with proper configuration
llm = Llama(
    model_path='C:/career growth/mistral-7b-instruct-v0.1.Q3_K_M.gguf',
    temperature=0.1,
    max_new_tokens=2000,
    context_window=3000,
    generate_kwargs={},
    model_kwargs={"n_gpu_layers": -1},
    verbose=True
)

async def extract_text_from_pdf(file_path):
    loader = PyPDFLoader(file_path)
    pages_content = ""
    async for page in loader.alazy_load():
        pages_content += page.page_content + "\n"
    return pages_content


def structure_with_ai(extracted_text):
    prompt = (
        f"Please structure the following resume text into JSON format with fields like user_email, language, Address, State, city, "
        f"bio, user_experience, education, phone, and linkedin. Only provide the structured JSON, without any additional explanations.\n\n"
        f"Text:\n{extracted_text}\n\nStructured JSON:"
    )
    response = llm(prompt)
    return response["choices"][0]["text"].strip()


def ocr_pdf_to_docx(pdf_file, docx_file):
    tesseract_path = "C:/career growth/Tesseract-OCR"
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

    images = convert_from_path(pdf_file)
    doc = Document()

    for page_num, img in enumerate(images, start=1):
        text = pytesseract.image_to_string(img, lang='eng')
        doc.add_paragraph(f"Page {page_num}")
        doc.add_paragraph(text)
    doc.save(docx_file)


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def correct_text_with_ai(text):
    prompt = (
        f"Correct the following text for grammar and spelling mistakes. "
        f"Return only the corrected text without any explanations:\n\n{text}"
    )
    response = llm(prompt)
    return response["choices"][0]["text"].strip()


def correct_text_for_specific_fields(structured_data, fields_to_check=["bio"]):
    field_corrections = {}
    for field in fields_to_check:
        if field in structured_data and structured_data[field]:
            original_text = structured_data[field]
            corrected_text = correct_text_with_ai(original_text)
            if original_text != corrected_text:
                field_corrections[field] = {
                    "original_text": original_text,
                    "corrected_text": corrected_text,
                }
    return field_corrections


def safe_temp_file(file_data, suffix=".pdf"):
    temp_dir = os.getenv("TEMP", "/tmp")  # Use TEMP environment variable or a safe fallback
    temp_file_path = tempfile.mkstemp(suffix=suffix, dir=temp_dir)[1]
    with open(temp_file_path, 'wb') as temp_file:
        temp_file.write(file_data)
    return temp_file_path


@app.post("/info")
async def read_resume(file: UploadFile = File(...)):
    temp_file_path = None
    docx_file_path = None

    try:
        # Save file to a safe temp directory
        file_data = await file.read()
        temp_file_path = safe_temp_file(file_data)

        if file.filename.endswith(".pdf"):
            extracted_text = await extract_text_from_pdf(temp_file_path)
        elif file.filename.endswith(".docx"):
            extracted_text = extract_text_from_docx(temp_file_path)
        else:
            return {"error": "Unsupported file format. Please upload a PDF or DOCX file."}

        if not extracted_text.strip():
            docx_file_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
            ocr_pdf_to_docx(temp_file_path, docx_file_path)
            extracted_text = extract_text_from_docx(docx_file_path)

        structured_data = structure_with_ai(extracted_text)
        structured_data = structured_data.replace('```json', '').replace('```', '').strip()
        json_data = json.loads(structured_data)

        corrections = correct_text_for_specific_fields(json_data)
        return JSONResponse(content={"structured_data": json_data, "corrections": corrections})

    except json.JSONDecodeError:
        return {"error": "AI-generated structure is not valid JSON."}

    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except PermissionError as e:
                print(f"Failed to delete temp file: {e}")
        if docx_file_path and os.path.exists(docx_file_path):
            try:
                os.remove(docx_file_path)
            except PermissionError as e:
                print(f"Failed to delete docx file: {e}")
